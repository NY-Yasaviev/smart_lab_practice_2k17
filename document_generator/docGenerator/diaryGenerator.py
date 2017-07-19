from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
diary = Document()
font = diary.styles['Normal'].font
font.name = 'Times New Roman'
font.size = Pt(14)
header = diary.add_paragraph('Дневник студента-практиканта\nХод выполнения практики\n').paragraph_format.alignment =  WD_PARAGRAPH_ALIGNMENT.CENTER
table = diary.add_table(rows=1,cols=4)
table.style = 'Table Grid'
#Создание и заполнение шапки
hat = table.rows[0]
hat.cells[0].add_paragraph('№').paragraph_format.alignment =  WD_PARAGRAPH_ALIGNMENT.CENTER
hat.cells[1].add_paragraph('Описание выполненной работы').paragraph_format.alignment =  WD_PARAGRAPH_ALIGNMENT.CENTER
hat.cells[2].add_paragraph('Дата').paragraph_format.alignment =  WD_PARAGRAPH_ALIGNMENT.CENTER
hat.cells[3].add_paragraph('Отметка\nруководителя').paragraph_format.alignment =  WD_PARAGRAPH_ALIGNMENT.CENTER
#Заполнение таблицы
for i in range(10):
    nextRow = table.add_row()
    nextRow.cells[0].add_paragraph(str(i+1)).paragraph_format.alignment =  WD_PARAGRAPH_ALIGNMENT.CENTER
    nextRow.cells[1].add_paragraph('some_text').paragraph_format.alignment =  WD_PARAGRAPH_ALIGNMENT.CENTER
    nextRow.cells[2].add_paragraph('some_date').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nextRow.cells[3].add_paragraph('some_mark').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#Форматирование ширины таблицы
table.autofit = False
for cell in table.columns[0].cells:
    cell.width = Inches(0.35)
for cell in table.columns[1].cells:
        cell.width = Inches(3)
for cell in table.columns[2].cells:
    cell.width = Inches(1.65)
for cell in table.columns[3].cells:
    cell.width = Inches(1.65)

diary.save('prepairDocx/Заполненный дневник.docx')