import random
import string

from django.shortcuts import render_to_response, render, redirect
from .models import *
from .forms import *
from .custom_decorators import is_deanery, is_student
from django.contrib.auth import get_user_model
from docxtpl import DocxTemplate
import os
import mimetypes
from django.http import StreamingHttpResponse
from wsgiref.util import FileWrapper
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


@is_deanery
def ind_list(request):
    title = "Индивидуальные задания"
    edu = IndividualTask.objects.filter(practice_type=IndividualTask.EDU)
    prod = IndividualTask.objects.filter(practice_type=IndividualTask.PROD)
    dip = IndividualTask.objects.filter(practice_type=IndividualTask.DIP)
    return render(request, 'deanery/ind_tasks.html', locals())


@is_deanery
def ind_edit(request, id):
    title = "Редактирование индивидуального задания"
    ind = IndividualTask.objects.get(pk=id)
    if request.POST:
        form = IndividualTaskForm(request.POST or None, instance=ind)
        if form.is_valid():
            form.save()
            return redirect("/ind_tasks/")
        else:
            return render(request, 'deanery/ind_task.html', locals())
    else:
        form = IndividualTaskForm(instance=ind)
        return render(request, 'deanery/ind_task.html', locals())


@is_deanery
def new_ind(request, type):
    if type == 'edu':
        choice = IndividualTask.EDU
    elif type == 'prod':
        choice = IndividualTask.PROD
    elif type == 'dip':
        choice = IndividualTask.DIP
    else:
        return redirect('/ind_tasks/')
    number = IndividualTask.objects.filter(practice_type=choice).count() + 1
    if request.POST:
        form = IndividualTaskForm(request.POST or None, initial={
            'practice_type': choice,
            'task_number': number
        })
        if form.is_valid():
            form.save()
            return redirect('/ind_tasks/')
        else:
            return render(request, 'deanery/addIndTask.html', {'form': form})
    else:
        form = IndividualTaskForm(initial={
            'practice_type': choice,
            'task_number': number
        })
        return render(request, 'deanery/addIndTask.html', {'form': form})


@is_deanery
def new_practice(request):
    title = "Добавление практики"
    if request.POST:
        form = PracticeForm(request.POST or None)
        if form.is_valid():
            form.save()
            return redirect("/practices/")
        else:
            return render(request, 'deanery/addPractice.html', {'form': form})
    else:
        form = PracticeForm()
        return render(request, 'deanery/addPractice.html', {'form': form})


@is_deanery
def edit_practice(request, id):
    title = "Редактирование практики"
    practice = Practice.objects.get(pk=id)
    # students_list = Student.objects.get(practice=practice)
    students_list = Student.objects.all().filter(practice=practice)
    if request.POST:
        form = PracticeForm(request.POST or None, instance=practice)
        if form.is_valid():
            form.save()
            return redirect("/practices/%s/" % id)
        else:
            return render(request, 'deanery/practice.html', {'form': form})
    else:
        form = PracticeForm(instance=practice)
        return render(request, 'deanery/practice.html', locals())


@is_deanery
def students(request):
    title = "Студенты"
    students_list = Student.objects.all()
    return render(request, 'deanery/students.html', locals())


@is_deanery
def new_student(request):
    title = "Добавление студента"
    if request.POST:
        form = StudentForm(request.POST or None)
        if form.is_valid():
            student = form.save(commit=False)
            login = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
            password = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
            user = get_user_model().objects.create_user(username=login, password=password)
            Group.objects.get(name='Students').user_set.add(user)
            student.login = login
            student.password = password
            student.user = user
            student.save()
            form.save_m2m()
            # creating docs
            for practice in student.practice.all():
                student_diary = Diary(student=student, practice=practice)
                student_diary.save()
            return redirect("/students/")
        else:
            return render(request, 'deanery/addStudent.html', {'form': form})
    else:
        form = StudentForm()
        return render(request, 'deanery/addStudent.html', {'form': form})


@is_deanery
def edit_student(request, id):
    title = "Редактирование студента"
    student = Student.objects.get(pk=id)
    if request.POST:
        form = StudentForm(request.POST or None, instance=student)
        if form.is_valid():
            form.save()
            return redirect("/students/")
    else:
        form = StudentForm(instance=student)
        return render(request, 'deanery/student.html', {'form': form}, locals())


@is_deanery
def practices(request):
    title = "Практики"
    practices_list = Practice.objects.all()
    return render(request, 'deanery/practices.html', locals())


@is_student
def profile(request):
    title = "Редактирование профиля"
    student = Student.objects.get(user=request.user)
    if request.POST:
        form = StudentForm(request.POST or None, instance=student)
        if form.is_valid():
            form.save()
            return redirect("/")
        else:
            return render(request, 'student/profile.html', locals())
    else:
        form = StudentForm(instance=student)
        return render(request, 'student/profile.html', locals())


@is_student
def practice_docs(request, id):
    return render(request, 'student/docs.html', locals())


@is_student
def new_diary_record(request, id):
    practice = Practice.objects.get(pk=id)
    student = Student.objects.get(user=request.user)
    d = Diary.objects.filter(student=student).get(practice=practice)
    if request.POST:
        form = DiaryRecordForm(request.POST or None)
        if form.is_valid():
            record = form.save(commit=False)
            record.diary = d
            record.save()
            return redirect('/practice_%s/diary/' % id)
        else:
            return render(request, 'student/addRecord.html', {'form': form})
    else:
        form = DiaryRecordForm()
        return render(request, 'student/addRecord.html', {'form': form})


@is_student
def diary_view(request, id):
    practice = Practice.objects.get(pk=id)
    diary = Diary.objects.get(practice=practice)
    records = DiaryRecord.objects.filter(diary=diary).order_by('date')
    return render(request, 'student/diaryView.html', locals())


@is_student
def diary(request, id):
    practice = Practice.objects.get(pk=id)
    diary = Diary.objects.get(practice=practice)
    records = DiaryRecord.objects.filter(diary=diary).order_by('date')
    return render(request, 'student/diary.html', locals())


@is_student
def edit_record(request, id, record_id):
    practice = Practice.objects.get(pk=id)
    diary = Diary.objects.get(practice=practice)
    record = DiaryRecord.objects.get(pk=record_id)
    if request.POST:
        form = DiaryRecordForm(request.POST or None, instance=record)
        if form.is_valid():
            form.save()
            return redirect('/practice_%s/diary/' % id)
        else:
            return render(request, 'student/record.html', {'form': form})
    else:
        form = DiaryRecordForm(instance=record)
        return render(request, 'student/record.html', locals())


@is_student
def diary_download(request, id):
    download(diary_save(request, id))


def diary_save(requset, id):
    practice = Practice.objects.get(pk=id)
    diary = Diary.objects.get(practice=practice)
    records = DiaryRecord.objects.filter(diary=diary)

    diary = Document()
    font = diary.styles['Normal'].font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    header = diary.add_paragraph(
        'Дневник студента-практиканта\nХод выполнения практики\n').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = diary.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    # Создание и заполнение шапки
    hat = table.rows[0]
    hat.cells[0].add_paragraph('№').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hat.cells[1].add_paragraph('Описание выполненной работы').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hat.cells[2].add_paragraph('Дата').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hat.cells[3].add_paragraph('Отметка\nруководителя').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Заполнение таблицы
    i = 1
    for record in records:
        nextRow = table.add_row()
        nextRow.cells[0].add_paragraph(str(i)).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nextRow.cells[1].add_paragraph('some_text').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nextRow.cells[2].add_paragraph('some_date').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nextRow.cells[3].add_paragraph('some_mark').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Форматирование ширины таблицы
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(0.35)
    for cell in table.columns[1].cells:
        cell.width = Inches(3)
    for cell in table.columns[2].cells:
        cell.width = Inches(1.65)
    for cell in table.columns[3].cells:
        cell.width = Inches(1.65)

    diary.save('prepairDocx/Заполненный дневник.docx')
    return 'docGenerator/prepairDocx/Заполненный дневник.docx'


@is_student
def individual(request, id):
    practice = Practice.objects.get(pk=id)
    return render(request, 'student/individual.html', locals())


# TODO FINISH THAT

@is_student
def edit_pass(request, id):
    practice = Practice.objects.get(pk=id)
    student = Student.objects.get(user=request.user)
    if request.POST:
        form = PassForm(request.POST or None, initial={
            'necessary_works': practice.necessary_works,
            'report': student.report,
            'review': student.review,
            'mark': student.mark
        })
        if form.is_valid():
            practice.necessary_works = form.cleaned_data['necessary_works']
            student.report = form.cleaned_data['report']
            student.review = form.cleaned_data['review']
            student.mark = form.cleaned_data['mark']
            return redirect("/practice_%s/pass/" % id)
        else:
            return render(request, 'student/pass.html')
    else:
        form = PassForm(initial={
            'necessary_works': practice.necessary_works,
            'report': student.report,
            'review': student.review,
            'mark': student.mark
        })
        return render(request, 'student/pass.html', locals())


@is_student
def pass_view(request, id):
    practice = Practice.objects.get(pk=id)
    return render(request,'student/passView.html',locals())


# TODO FINISH THAT

@is_student
def report_view(request, id):
    practice = Practice.objects.get(pk=id)
    if practice.type == 'Учебная':
        type = 'УЧЕБНОЙ'
    else:
        type = 'ПРОИЗВОДСТВЕННОЙ'
    student = Student.objects.get(user=request.user)
    return render(request, 'student/reportView.html', locals())


@is_student
def report_download(request, id):
    download(report_save(request, id))


def report_save(request, id):
    s = Student.objects.get(user=request.user)
    p = Practice.objects.get(pk=id)

    report = DocxTemplate("docGenerator/templates/report.docx")

    to_render = {'typePractice': "",
                 'orgName': p.company,
                 'practLeader': p.teacher,
                 'studGroup': s.group,
                 'compLeader': p.chief,
                 'studFio': s.name
                 }
    if p.type == p.PROD or p.type == p.DIP:
        to_render['typePractice'] = "ПРОИЗВОДСТВЕННАЯ"
    else:
        to_render['typePractice'] = "УЧЕБНАЯ"

    report.render(to_render)

    report.save("prepairDocx/Заполненный отчет.docx")
    return "docGenerator/prepairDocx/Заполненный отчет.docx"


@is_student
def individual_view(request, id):
    practice = Practice.objects.get(pk=id)
    if practice.type == 'Учебная':
        type = 'учебную'
    elif practice.type == 'Производственная':
        type = 'производственную'
    else:
        type = 'преддипломную'
    student = Student.objects.get(user=request.user)
    inds = IndividualTask.objects.filter(practice_type=practice.type)
    return render(request, 'student/indView.html', locals())


@is_student
def individual_download(request, id):
    download(individual_save(request, id))


def individual_save(request, id):
    s = Student.objects.get(user=request.user)
    p = Practice.objects.get(pk=id)
    inds = IndividualTask.objects.filter(practice_type=p.type)
    filename = 'templates/indPr' + inds.count() + '.docx'
    doc = DocxTemplate(filename)
    to_render = {'yearF': "2016",
                 'yearT': "2017",
                 'profil': s.edu_profile,
                 'entityName': s.company,
                 'studFio': s.name,
                 'studCours': s.course,
                 'studGroup': s.group,
                 'dateF': p.date_from,
                 'dateT': p.date_to,
                 'tutorKfuFio': p.teacher,
                 'tutorKfuStatus': '',
                 'tutorEntityFio': p.chief,
                 'tutorKfuStatus': '',
                 'typePractice': p.type.upper()
                 }
    i = 1
    for ind in inds.order_by('id'):
        if i > 5:
            i = len(inds)

        to_render['task' + i] = ind.desc
        to_render['dateF' + i] = ind.dateFrom
        to_render['dateT' + i] = ind.dateFrom
        i += 1
    doc.render(to_render)
    doc.save('prepairDocx/Индивидуальное задание.docx')
    return 'docGenerator/prepairDocx/Индивидуальное задание.docx'


@is_student
def pass_download(request, id):
    download(pass_save(request, id))
    # may be redirect(...)


def pass_save(request, id):
    # preparation
    s = Student.objecst.get(user=request.user)
    p = Practice.objecst.get(pk=id)
    # do IT
    doc = DocxTemplate("templates/permit.docx")
    changeTag = {'studCours': s.course,
                 'studGroup': s.group,
                 'departName': "Высшая школа ИТИС КФУ",
                 'contNumber': "",
                 'contDate': "",
                 'entityName': p.company,
                 'entityAdress': p.address,
                 'practType': p.type,
                 'practDateT1': p.date_from,
                 'practDateF1': p.date_to,
                 'practDateT2': "",
                 'practDateF2': "",
                 'practDateT3': "",
                 'practDateF3': "",
                 'corollary': s.review,
                 'mark': s.mark,
                 'spec': s.edu_profile,
                 'profil': "",
                 'skill': s.status,
                 'needActivType': p.necessary_works,
                 'showUp': p.date_from,
                 'tutor': p.chief,
                 'adminReview': s.report,
                 'departure': p.date_to}
    doc.render(changeTag)
    doc.save("prepairDocx/Заполненная путевка.docx")
    return 'docGenerator/prepairDocx/Заполненная путевка.docx'


def download(file):
    the_file = file
    filename = os.path.basename(the_file)
    chunk_size = 8192
    response = StreamingHttpResponse(FileWrapper(open(the_file, 'rb'), chunk_size),
                                     content_type=mimetypes.guess_type(the_file)[0])
    response['Content-Length'] = os.path.getsize(the_file)
    response['Content-Disposition'] = "attachment; filename=%s" % filename
    return response


def report(request, id):
    practice = Practice.objects.get(pk=id)
    return render(request, 'student/report.html', locals())
